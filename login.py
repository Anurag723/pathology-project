import customtkinter as ctk
import tkinter as tk
import sqlite3
from tkcalendar import DateEntry
from tkinter import messagebox,ttk
from PIL import Image
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx import Document
from docx.oxml.ns import qn
import fitz
import os
import shutil
import subprocess
from docxtpl import DocxTemplate
from docx.shared import Inches
from docx.enum.text import WD_BREAK

try:
    import win32print
    import win32api
except ImportError:
    win32print = None


def entry():
    root.withdraw()

    index = ctk.CTkToplevel(root)
    index.state("zoomed")

    def exit_app():
        confirm = messagebox.askyesno("Exit", "Are you sure you want to exit?")
        if confirm:
            root.destroy()

    def fetch_test_names():
        """Fetch test names from the SQLite database."""
        conn = sqlite3.connect("test_and_doctor.db")
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM test")  # Fetch test names
        names = [row[0] for row in cursor.fetchall()]
        conn.close()
        return names

    def new_entry(parent):
        try:
            new = ctk.CTkToplevel(parent)
            new.title("New Report")
            new.geometry("1100x650")
            # new.state("zoomed")
            new.minsize(1100,650)

            frame = ctk.CTkFrame(new, fg_color="#9bfcf9")
            frame2 = ctk.CTkFrame(new, fg_color="#9bfcf9")

            frame.place(relx=0.0, rely=0, relwidth=0.51, relheight=1)  # Left half
            frame2.place(relx=0.515, rely=0, relwidth=0.49, relheight=1)  # Right half

            # add data to the data base---------------------------------------
            def save():
                
                name = abb_var.get()+patient_entry.get()
                ref = ref_entry.get().strip()
                mobile = mob_entry.get().strip()
                entrydate = entry.get()
                deliverydate = delivery.get()
                gender = gen_dropdown.get().strip()
                age = age_entry.get().strip()+age_dropdown.get()
                refby = refby_entry.get().strip()
                test = ",".join(allchoice)
                paid = paid_entry.get()
                remaining = blnc_entry.get()
                file = patient_entry.get().strip()+ref

                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()

                curs.execute('''INSERT INTO patient(name,ref_no,mobile,entry,delivery,gender,age,ref_by,test,paid,remaining,file)VALUES(?,?,?,?,?,?,?,?,?,?,?,?)''',(name,ref,mobile,entrydate,deliverydate,gender,age,refby,test,paid,remaining,file))

                conn.commit()
                conn.close()

                messagebox.showinfo("Done!","Data added successfully!!!")

            def get_doctor_names():
                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()
                curs.execute("SELECT name FROM doctor")  # Adjust column name if needed
                doctors = [row[0] for row in curs.fetchall()]
                conn.close()
                return doctors

            # Patient Details Labels-------------------------------------------
            patient_label = ctk.CTkLabel(frame, text="Patient:", text_color="black")
            patient_label.grid(row=0, column=0, padx=10, pady=5, sticky="w")
            abb_var = tk.StringVar()
            abb_dropdown = ttk.Combobox(frame, textvariable=abb_var, values=["Mr.", "Mrs.", "Ms.", "Dr."], width=3, state="readonly")
            abb_dropdown.grid(row=0, column=1, padx=5, pady=5)
            abb_dropdown.current(0)
            patient_entry = ctk.CTkEntry(frame, fg_color="white", text_color="black")
            patient_entry.grid(row=0, column=2, pady=5)

            # ref no.-------------------------------------------
            ref_label = ctk.CTkLabel(frame, text="Ref No.", text_color="black")
            ref_label.grid(row=0, column=3, padx=10, pady=5, sticky="w")
            ref_entry = ctk.CTkEntry(frame, fg_color="white", text_color="black", width=50)
            ref_entry.grid(row=0, column=4, padx=10, pady=5, sticky='w')

            # mobile no:-------------------------------------------
            mob_label = ctk.CTkLabel(frame, text="Mobile No.:", text_color="black")
            mob_label.grid(row=1, column=0, padx=10, pady=3, sticky="w")
            mob_entry = ctk.CTkEntry(frame, fg_color="white", text_color="black")
            mob_entry.grid(row=1, column=1, columnspan=2, padx=10, pady=3, sticky="w")

            # enter calender-------------------------------------------
            enter_label = ctk.CTkLabel(frame, text="Entry Date:", text_color="black")
            enter_label.grid(row=1, column=3, padx=5, pady=3, sticky='w')
            entry = DateEntry(frame, width=12, background="black", foreground="white", borderwidth=2,date_pattern = "yyyy/mm/dd", state="readonly")
            entry.grid(row=1, column=4, columnspan=1, padx=5, pady=3, sticky='w')

            # address-------------------------------------------
            address_label = ctk.CTkLabel(frame, text="Address:", text_color="black")
            address_label.grid(row=2, column=0, padx=10, pady=3, sticky="w")
            address_entry = ctk.CTkEntry(frame, fg_color="white", text_color="black", width=180)
            address_entry.grid(row=2, column=1, columnspan=4, padx=10, pady=3, sticky="w")

            # delivery calender-------------------------------------------
            delivery_label = ctk.CTkLabel(frame, text="Delivery Date:", text_color="black")
            delivery_label.grid(row=2, column=3, padx=5, pady=3, sticky='w')
            delivery = DateEntry(frame, width=12, background="green", foreground="white", borderwidth=2,date_pattern = "yyyy/mm/dd", state="readonly")
            delivery.grid(row=2, column=4, columnspan=1, padx=5, pady=3, sticky='w')

            # gender-------------------------------------------
            gender_label = ctk.CTkLabel(frame, text="Gender:", text_color="black")
            gender_label.grid(row=3, column=0, padx=10, pady=5, sticky="w")
            gen_var = tk.StringVar()
            gen_dropdown = ttk.Combobox(frame, textvariable=gen_var, values=["Male", "Female", "Other"], width=6, state="readonly")
            gen_dropdown.grid(row=3, column=1, padx=5, pady=5)
            gen_dropdown.current(0)

            # age-------------------------------------------
            ctk.CTkLabel(frame, text="Age:", text_color="black").grid(row=3, column=2, padx=(10,3), pady=5, sticky='e')
            age_entry = ctk.CTkEntry(frame, width=50, fg_color="white", text_color="black")
            age_entry.grid(row=3, column=3, pady=5, sticky='e')
            age_var = tk.StringVar()
            age_dropdown = ttk.Combobox(frame, textvariable=age_var, values=["Yrs", "Months", "Days"], width=6, state="readonly")
            age_dropdown.grid(row=3, column=4, pady=5, sticky='w')
            age_dropdown.current(0)

            # ref by-------------------------------------------
            doctor_names = get_doctor_names()
            refby_label = ctk.CTkLabel(frame, text="Ref by:", text_color="black")
            refby_label.grid(row=4, column=0, padx=10, pady=3, sticky="w")
            refby_entry = ttk.Combobox(frame, values=doctor_names, width=30)
            refby_entry.grid(row=4, column=1, columnspan=4, padx=10, pady=3, sticky="w")

            # for test details--------------------------------
            detframe = ctk.CTkFrame(frame, fg_color='white')
            detframe.place(relx=0.02, rely=0.35, relwidth=0.96, relheight=0.3)

            # Label to show selected tests
            selected_tests = set()  # Set to store selected test names
            latest_choice = None  # Variable to store the last selected choice

            detframe = ctk.CTkFrame(frame, fg_color='white')
            detframe.place(relx=0.02, rely=0.35, relwidth=0.96, relheight=0.3)

            selected_test_label = ctk.CTkLabel(detframe, text="No Test Selected", fg_color="white",text_color="violet", width=200, height=30, font=("Arial", 22, "bold"))
            selected_test_label.pack(pady=30)

            

            framer = tk.Frame(frame2, bg="white", padx=10, pady=10)
            framer.place(relx=0.01, rely=0.01, relwidth=0.97, relheight=0.55)


            # ---- GLOBAL VARIABLES ----
            stored_data = {}  # Store user-edited values

            def fetch_table_data(table_name):
                """Fetch column1, column3, and column4 from the selected database table."""
                if not table_name:
                    return []  # Return empty if no table is selected
                with sqlite3.connect("test_and_doctor.db") as conn:
                    return conn.execute(f"SELECT type, value, unit FROM {table_name}").fetchall()

            def create_table(frame, table_name):
                """Creates a table inside the given frame with editable Column 2 and a scrollbar."""
                for widget in frame.winfo_children():  # Clear existing treeview
                    widget.destroy()

                nonlocal stored_data
                stored_data = {}

                tree_frame = tk.Frame(frame)
                tree_frame.pack(fill="both", expand=True, padx=5, pady=5)

                columns = ("Column1", "Column2", "Column3", "Column4")
                tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=10)

                for col, text in zip(columns, ["Test Name", "Your Report", "Healthy Report", "Unit"]):
                    tree.heading(col, text=text)
                    tree.column(col, width=150, anchor="center")

                scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
                tree.configure(yscrollcommand=scrollbar.set)

                tree.pack(side="left", fill="both", expand=True)
                scrollbar.pack(side="right", fill="y")

                for row in fetch_table_data(table_name):
                    item_id = tree.insert("", "end", values=(row[0], "", row[1], row[2]))
                    stored_data[item_id] = [row[0], "", row[1], row[2]]

                entry = ttk.Entry(frame)

                def on_double_click(event):
                    """Opens entry box when double-clicking a cell in Column 2."""
                    item_id = tree.selection()[0] if tree.selection() else None
                    if item_id and tree.identify_column(event.x) == "#2":
                        x, y, width, height = tree.bbox(item_id, 1)
                        entry.place(x=x + 5, y=y + height // 2, width=width, height=height)
                        entry.focus()
                        entry.delete(0, "end")
                        entry.insert(0, tree.item(item_id, "values")[1])

                        entry.bind("<Return>", lambda e: save_entry(item_id))

                def save_entry(item_id):
                    """Saves user input to Column 2."""
                    stored_data[item_id][1] = entry.get()
                    tree.item(item_id, values=stored_data[item_id])
                    entry.place_forget()

                tree.bind("<Double-1>", on_double_click)

                # Print Data Button
                # ctk.CTkButton(frame, text="Print Data", command=lambda: print("\nStored Data:", *stored_data.values()), fg_color="lightblue", text_color="black").pack(pady=10)

                # Generate Report Button
                ctk.CTkButton(
                    frame, text="Generate Report", command=lambda: generate_pathology_report(table_name),
                    fg_color="green", text_color="white"
                ).pack(pady=10)

                return tree
           
            
            def generate_pathology_report(selected_tests):
                """
                Appends new reports to an existing file if it exists, 
                otherwise, creates a new pathology report.
                - Each test appears on a new page.
                - Static details from the template (Lab Name, Address, etc.) are retained on every page.
                - Patient details remain the same.
                - The heading and table data change for each test.
                - The template image is preserved in the generated output.
                """

                output_path = patient_entry.get().strip() + ref_entry.get().strip() + ".docx"
                template_path = "invoice.docx"

                # Example patient details (replace with actual input fields)
                patient_info = {
                    "name": f"{abb_var.get()} {patient_entry.get()}",
                    "age": f"{age_entry.get()} {age_dropdown.get()}",
                    "gender": gen_dropdown.get(),
                    "mno": mob_entry.get(),
                    "address": address_entry.get(),
                    "date": entry.get(),
                    "latest_test": ", ".join(selected_tests)  # Save as a whole string
                }

                # If the report exists, load it; otherwise, start a blank report
                if os.path.exists(output_path):
                    print("📄 Existing report found. Appending new tests...")
                    final_doc = Document(output_path)
                else:
                    print("🆕 Creating a new report...")
                    final_doc = Document()  # Create a blank document

                # Fetch table data once for all tests
                table_data = [
                    {"Column1": row[0], "Column2": row[1], "Column3": row[2], "Column4": row[3]}
                    for row in stored_data.values()
                ]
                patient_info["report_table"] = table_data

                # Update only the latest test name for each iteration
                patient_info["latest_test"] = latest_choice

                # Load the template (preserving images & static details)
                doc = DocxTemplate(template_path)
                doc.render(patient_info)

                # Save the rendered template to a temporary file
                temp_output = "temp_report.docx"
                doc.save(temp_output)

                # Load the rendered document
                temp_doc = Document(temp_output)

                # ✅ Force a Page Break BEFORE adding new content
                if final_doc.paragraphs:  # Avoid adding a page break if it's the first page
                    page_break = OxmlElement("w:br")
                    page_break.set(qn("w:type"), "page")
                    final_doc.element.body.append(page_break)

                # ✅ Append only the STRUCTURE & CONTENT of the template (not the template itself)
                for element in temp_doc.element.body:
                    final_doc.element.body.append(element)

                # Save the final report
                final_doc.save(output_path)

                # Clean up temporary file
                if os.path.exists(temp_output):
                    os.remove(temp_output)

                print(f"✅ Report updated: {output_path}")


            stotal = 0  # Initialize stotal outside the functions
            allchoice = set()

            def fetch_test_price(test_name):
                """Fetch price of the selected test."""
                if not test_name:
                    return "0.00"  # Default value if no test is selected
                with sqlite3.connect("test_and_doctor.db") as conn:
                    result = conn.execute("SELECT price FROM test WHERE name = ?", (test_name,)).fetchone()
                    return result[0] if result else "0.00"
                
            def update_balance():
                """Calculates balance as stotal - discount - paid amount and updates blnc_entry."""
                nonlocal stotal

                try:
                    discount = float(dis_entry.get()) if dis_entry.get() else 0  # Get discount safely
                except ValueError:
                    discount = 0  # Handle invalid input

                try:
                    paid_amount = float(paid_entry.get()) if paid_entry.get() else 0  # Get paid amount safely
                except ValueError:
                    paid_amount = 0  # Handle invalid input

                balance = stotal - discount - paid_amount  # Calculate balance

                # Update balance entry for new test
                blnc_entry.configure(state="normal")
                blnc_entry.delete(0, "end")
                blnc_entry.insert(0, str(balance))
                blnc_entry.configure(state="readonly")

            def update_label(choice):
                """Updates the table when a new test is selected and sets the price."""
                nonlocal latest_choice, stotal, allchoice  # Use nonlocal if inside another function

                if choice and choice != "Select Test":
                    # Check if the test is already in allchoice
                    if choice in allchoice:
                        messagebox.showwarning("already done",f"{choice} is already selected. No updates made.")
                        return  # Exit the function early if the choice is already in allchoice

                    selected_tests.add(choice)
                    latest_choice = choice  # Store latest selected choice
                    allchoice.add(latest_choice)
                    selected_test_label.configure(text=", ".join(selected_tests))
                    print("Latest Choice:", latest_choice)  # Automatically print latest selection

                    # Fetch price and update sbttl_entry
                    price = fetch_test_price(latest_choice)
                    price = float(price) if price else 0  # Convert price to float

                    # Update stotal
                    stotal += price

                    sbttl_entry.configure(state="normal")  # Enable editing
                    sbttl_entry.delete(0, "end")  # Clear existing value
                    sbttl_entry.insert(0, str(price))  # Set new price
                    sbttl_entry.configure(state="readonly")  # Make it readonly again

                    ttl_entry.configure(state="normal")
                    ttl_entry.delete(0, "end")
                    ttl_entry.insert(0, stotal)
                    ttl_entry.configure(state="readonly")

                    update_balance()

                    create_table(framer, latest_choice)  # Refresh the table with new data

            def add_update_blnc():
                paid_amount = float(paid_entry.get())
                ref_number = ref_entry.get().strip()

                conn = sqlite3.connect("test_and_doctor.db")  # Replace with actual database path
                cursor = conn.cursor()

                # Fetch the current value from the 'remaining' column where ref_no matches
                cursor.execute("SELECT remaining FROM patient WHERE ref_no = ?", (ref_number,))
                result = cursor.fetchone()

                cursor.execute("SELECT paid FROM patient WHERE ref_no = ?", (ref_number,))
                resultpd = cursor.fetchone()

                if result:
                    remaining_amount = float(result[0])  # Convert the current value to float
                    pdamt = float(resultpd[0])
                    new_remaining = max(remaining_amount - paid_amount, 0)  # Ensure it doesn't go negative
                    new_paid_amount = pdamt + paid_amount

                    # Update the database with the new remaining amount
                    cursor.execute("UPDATE patient SET paid=? , remaining = ? WHERE ref_no = ?", (new_paid_amount,new_remaining, ref_number))
                    conn.commit()
                    messagebox.showinfo("Done",f"✅ Updated remaining amount: {new_remaining} and total paid {new_paid_amount}")
                    
                else:
                    messagebox.showinfo("❌ Reference number not found!")

                conn.close()



            # Fetch test names
            test_name = fetch_test_names()

            # Test selection UI
            ctk.CTkLabel(frame, text="Test", text_color="indigo").grid(row=5, column=0, padx=10, pady=3)

            test_entry = ctk.CTkComboBox(frame, values=test_name, width=350, fg_color="#f8fc9f", text_color="dark blue", command=update_label, state="readonly")

            test_entry.set("Select Test")
            test_entry.grid(row=5, column=1, columnspan=6, padx=3, pady=3, sticky="w")


            # amount frame-------------------------
            amtframe = ctk.CTkFrame(frame,fg_color="#9bfcf9")
            amtframe.place(relx=0, rely=0.67, relwidth=0.6, relheight= 0.23)

            ctk.CTkLabel(amtframe, text="Out S.total", text_color='black').grid(row=0, column=0, padx=10, pady=5)
            outtotal_entry = ctk.CTkEntry(amtframe, width=80, fg_color="#EDE9E8", state='readonly')
            outtotal_entry.grid(row=0, column=1, padx=10, pady=5, sticky="w")

            ctk.CTkLabel(amtframe, text="Mode:", text_color="black").grid(row=1, column=0, padx=10, pady=3)

            mode_var = tk.StringVar()
            mode_dropdown = ttk.Combobox(amtframe, textvariable=mode_var, values=["UPI.", "Cash", "Card", "Other"], width=5, state="readonly")
            mode_dropdown.grid(row=1, column=1, padx=1, pady=3, sticky="w")
            mode_dropdown.current(0)

            ctk.CTkLabel(amtframe, text="status", text_color="black").grid(row=1, column=3, padx=1, pady=3)
            
            status_var = tk.StringVar()
            status_dropdown = ttk.Combobox(amtframe, textvariable=status_var, values=["pending", "clear"], width=5, state="readonly")
            status_dropdown.grid(row=1, column=4, padx=0, pady=3, sticky="w")
            status_dropdown.current(0)

            ctk.CTkLabel(amtframe, text="Note.", text_color="black").grid(row=2, column=0, padx=10, pady=3)

            note_entry = ctk.CTkEntry(amtframe, width=200, fg_color="#D5FAF2", text_color="black")
            note_entry.grid(row=2, column=1, padx=5, pady=3, columnspan=7, sticky="w")

            # payment-----------------------------------------------
            payframe=ctk.CTkFrame(frame, fg_color="#9bfcf9")
            payframe.place(relx=0.63, rely=0.67, relwidth=0.38, relheight=0.23)

            ctk.CTkLabel(payframe, text="Subtotal", text_color="black").grid(row=0, column=0, padx=3, pady=5)
            sbttl_entry=ctk.CTkEntry(payframe,text_color="black", width=100, fg_color="#E3E7E6", state="readonly")
            sbttl_entry.grid(row=0, column=1, padx=7, pady=1)

            ctk.CTkLabel(payframe, text="Discount", text_color="black").grid(row=1, column=0, padx=3, pady=5)
            dis_entry=ctk.CTkEntry(payframe,text_color="black", width=100, fg_color="#E3E7E6")
            dis_entry.grid(row=1, column=1, padx=7, pady=1)
            dis_entry.bind("<KeyRelease>", lambda e: update_balance())
            
            ctk.CTkLabel(payframe, text="Total", text_color="black").grid(row=2, column=0, padx=3, pady=5)
            ttl_entry=ctk.CTkEntry(payframe,text_color="black", width=100, fg_color="#E3E7E6", state="readonly")
            ttl_entry.grid(row=2, column=1, padx=7, pady=1)
            
            ctk.CTkLabel(amtframe, text="Paid Amount", text_color="black").grid(row=3, column=0, padx=3, pady=5)
            paid_entry=ctk.CTkEntry(amtframe,text_color="black", width=100, fg_color="#E3E7E6")
            paid_entry.grid(row=3, column=1, padx=7, pady=1)
            paid_entry.bind("<KeyRelease>", lambda e: update_balance())

            
            ctk.CTkLabel(payframe, text="Balance", text_color="black").grid(row=3, column=0, padx=3, pady=5)
            blnc_entry=ctk.CTkEntry(payframe,text_color="black", width=100, fg_color="#E3E7E6", state="readonly")
            blnc_entry.grid(row=3, column=1, padx=7, pady=1)

            # button--------------------------------------
            btnfrm=ctk.CTkFrame(frame, fg_color="#9bfcf9")
            btnfrm.place(relx=0, rely=0.9, relwidth=1, relheight=0.1)
            
            update = ctk.CTkButton(btnfrm, width=100, text="Update Bal.", fg_color="#E49FF8", text_color="black", command=add_update_blnc)
            update.grid(row=0, column=0, padx=20, pady=5)

            add_new = ctk.CTkButton(btnfrm, width=100, text="Add New", fg_color="#FFECA1", text_color="black")
            add_new.grid(row=0, column=1, padx=20, pady=5)
            
        #--------------------heading frame---------------------------------------------------
         
            head = ctk.CTkFrame(frame2, fg_color="#9bfcf9")
            head.place(relx=0, rely=0.56, relwidth=1, relheight=0.44)

            ctk.CTkLabel(head, text="Heading", text_color="black").grid(row=0, column=0, padx=10, pady=(5,0), sticky="w")
            head_entry=ctk.CTkEntry(head, width=200, fg_color="white", text_color= "black")
            head_entry.grid(row=1, column=0,padx=10, pady=5, sticky="w")

            ctk.CTkLabel(head, text="Reporting Date", text_color="black").grid(row=0, column=1, padx=10, pady=(5,0))
            
            rpt_date = DateEntry(head, width=12, background="#FE022A", foreground="#FFE4E9", borderwidth=2, date_pattern="yyyy/mm/dd", state="readonly")
            rpt_date.grid(row=1, column=1, padx=10, pady=5)
            
            ctk.CTkLabel(head, text="Method", text_color="black").grid(row=2, column=0, padx=10, pady=(5,0), sticky="w")
            mthd_entry = ctk.CTkEntry(head, width=400, text_color="black", fg_color="white")
            mthd_entry.grid(row=3, column=0, columnspan=3, padx=10, pady=5)

            ctk.CTkLabel(head, text="Comments", text_color="black").grid(row=4, column=0, padx=10, pady=(5,0), sticky="w")

            cmnt_entry = ctk.CTkEntry(head, width=400, height=56, fg_color="white", text_color="black")
            cmnt_entry.grid(row=5, column=0, padx=10, pady=5, columnspan=5)
            
            ctk.CTkButton(head, text="Save", fg_color="#08B5BC", text_color="black", command=save).grid(row=6, column=0, padx=8, pady=5)

            ctk.CTkButton(head, text="Print", width=100, fg_color="white", text_color="black", command=lambda:print_entry(new)).grid(row=6, column=1, padx=8, pady=5, sticky="w")

            ctk.CTkButton(head, text="close",width=70,  fg_color="red", text_color="black", command=new.destroy).grid(row=6, column=2, padx=8, pady=5, sticky="w")
            new.grab_set()

        except Exception as e:
            messagebox.showerror("Error",e)

    def view_entry(parent):
        view = ctk.CTkToplevel(parent)
        view.title("View Report")
        view.geometry("950x650")
        frame = ctk.CTkFrame(view, fg_color="#a4fcca")
        frame.place(relx=0 ,rely=0, relwidth=1, relheight=1)

        # -------------doctor data-----------------------------
        def doctor_det():
            # Function to fetch data from the database
            def fetch_data():
                conn = sqlite3.connect('test_and_doctor.db')
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM doctor")
                rows = cursor.fetchall()
                conn.close()
                return rows

            # Function to populate the treeview with data
            def populate_treeview(treeview, data):
                for row in data:
                    treeview.insert("", "end", values=row)

            # **Clear the previous content from frame_det**
            for widget in frame_det.winfo_children():
                widget.destroy()

            # Create a style
            style = ttk.Style()
            style.configure("Treeview", font=("Helvetica", 14))  # Font size for the treeview content
            style.configure("Treeview.Heading", font=("Helvetica", 16, "bold"))  # Font size for the headings
            
            # Create the treeview
            columns = ("Name", "Specialization", "Location")
            treeview = ttk.Treeview(frame_det, columns=columns, show="headings")

            # Define the headings and column widths
            for col in columns:
                treeview.heading(col, text=col)
                treeview.column(col, width=100)

            # Fetch data from the database and populate the treeview
            data = fetch_data()
            populate_treeview(treeview, data)

            # Add the treeview to the frame with a scrollbar
            scrollbar = ttk.Scrollbar(frame_det, orient="vertical", command=treeview.yview)
            treeview.configure(yscrollcommand=scrollbar.set)
            treeview.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
        
        def test_det():
            """Displays test details from the database in a treeview."""
            
            # Function to fetch data from the database
            def fetch_data():
                conn = sqlite3.connect('test_and_doctor.db')
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM test")  # Ensure correct column names
                rows = cursor.fetchall()
                conn.close()
                return rows

            # Function to populate the treeview with data
            def populate_treeview(treeview, data):
                treeview.delete(*treeview.get_children())  # Clear previous data
                for row in data:
                    treeview.insert("", "end", values=row)

            # **Clear the previous content from frame_det**
            for widget in frame_det.winfo_children():
                widget.destroy()

            # Create a frame inside frame_det to hold the treeview and scrollbars
            tree_frame = ttk.Frame(frame_det)
            tree_frame.pack(fill="both", expand=True)

            # Create the treeview
            columns = ("name", "price")  # Match with database column names
            treeview = ttk.Treeview(tree_frame, columns=columns, show="headings")

            # Define the headings and column widths
            treeview.heading("name", text="Test Name")
            treeview.heading("price", text="Price")
            treeview.column("name", width=200)  # Increased width for better visibility
            treeview.column("price", width=100)

            # **Increase font size for better readability**
            style = ttk.Style()
            style.configure("Treeview", font=("Arial", 14))  # Adjust font size
            style.configure("Treeview.Heading", font=("Arial", 16, "bold"))  # Adjust header font

            # Fetch data from the database and populate the treeview
            data = fetch_data()
            populate_treeview(treeview, data)

            # **Add Vertical and Horizontal Scrollbars**
            y_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=treeview.yview)
            x_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=treeview.xview)

            treeview.configure(yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set)

            # **Pack the treeview and scrollbars properly**
            y_scrollbar.pack(side="right", fill="y")
            x_scrollbar.pack(side="bottom", fill="x")
            treeview.pack(side="left", fill="both", expand=True)

        def patient_det():
            try:
                # Function to fetch data from the database
                def fetch_data():
                    conn = sqlite3.connect('test_and_doctor.db')
                    cursor = conn.cursor()
                    cursor.execute("SELECT * FROM patient")
                    rows = cursor.fetchall()
                    conn.close()
                    return rows

                # Function to populate the treeview with data
                def populate_treeview(treeview, data):
                    for row in data:
                        treeview.insert("", "end", values=row)

                # **Clear the previous content from frame_det**
                for widget in frame_det.winfo_children():
                    widget.destroy()

                # Create a style
                style = ttk.Style()
                style.configure("Treeview", font=("Helvetica", 14))  # Font size for the treeview content
                style.configure("Treeview.Heading", font=("Helvetica", 16, "bold"))  # Font size for the headings

                # Create a frame inside frame_det to hold the treeview and scrollbars
                tree_frame = ttk.Frame(frame_det)
                tree_frame.pack(fill="both", expand=True)

                # Create the treeview
                columns = ("Name", "Ref No.", "Mobile", "Entry Date", "Delivery Date", "Gender",
                        "Age", "Ref by", "Test", "Paid", "Remaining", "File")

                treeview = ttk.Treeview(tree_frame, columns=columns, show="headings")

                # Define the headings and column widths
                for col in columns:
                    treeview.heading(col, text=col)
                    treeview.column(col, width=150)  # Adjust width as needed

                # Fetch data from the database and populate the treeview
                data = fetch_data()
                populate_treeview(treeview, data)

                # **Add Scrollbars**
                y_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=treeview.yview)
                x_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=treeview.xview)
                
                treeview.configure(yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set)

                # **Pack everything**
                y_scrollbar.pack(side="right", fill="y")
                x_scrollbar.pack(side="bottom", fill="x")
                treeview.pack(side="left", fill="both", expand=True)

            except Exception as e:
                messagebox.showerror("Error!",e)

        # buttons----------------------------
        ctk.CTkButton(frame, text="Doctor Record", corner_radius=3, fg_color="red", text_color="white", command=doctor_det).grid(row=0, column=0, padx=50, pady=10)
        ctk.CTkButton(frame, text="Test Record", corner_radius=3, fg_color="blue", text_color="white", command=test_det).grid(row=0, column=1, padx=50, pady=10)
        ctk.CTkButton(frame, text="Patient Record", corner_radius=3, fg_color="green", text_color="white", command=patient_det).grid(row=0, column=2, padx=50, pady=10)

        # data frame--------------------------
        frame_det = ctk.CTkFrame(frame, fg_color="grey")
        frame_det.place(relx=0.01, rely=0.13, relwidth=0.98, relheight=0.5)
        view.grab_set()
        
    def print_entry(parent):
        """Opens a Print Preview Window with multiple pages"""
        print_win = ctk.CTkToplevel(parent, fg_color="#e0f0ef")
        print_win.title("Print Report")
        print_win.geometry("950x650")

        frame = ctk.CTkFrame(print_win, fg_color="#92c4bf")
        frame.place(relx=0, rely=0, relwidth=0.3, relheight=1)

        # PDF Variables
        doc = None
        num_pages = 0
        current_page = [0]

        # Image Display
        image_label = ctk.CTkLabel(print_win, text="")
        image_label.place(relx=0.4, rely=0.1)

        def update_preview():
            """Update the preview with the current page"""
            if doc and num_pages > 0:
                page = doc[current_page[0]]
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img.thumbnail((300, 450))  # Resize image to fit

                # Convert to CTkImage (Fixes the Warning)
                ctk_img = ctk.CTkImage(light_image=img, size=(300, 450))

                # Update Label with CTkImage
                image_label.configure(image=ctk_img, text="")
                image_label.image = ctk_img  # Keep reference

                # Enable/Disable buttons based on page number
                btn_prev.configure(state="normal" if current_page[0] > 0 else "disabled")
                btn_next.configure(state="normal" if current_page[0] < num_pages - 1 else "disabled")

        def open_file():
            """Open the PDF file specified in the search box"""
            nonlocal doc, num_pages, current_page
            filename = search.get().strip()+".pdf"

            if not filename:
                messagebox.showwarning("Warning", "Please enter a file name.")
                return

            if not os.path.exists(filename):
                messagebox.showerror("Error", "File not found. Please check the filename.")
                return

            try:
                doc = fitz.open(filename)  # Open PDF
                num_pages = len(doc)
                current_page[0] = 0
                update_preview()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open file:\n{e}")

        def next_page():
            """Go to the next page"""
            if doc and current_page[0] < num_pages - 1:
                current_page[0] += 1
                update_preview()

        def prev_page():
            """Go to the previous page"""
            if doc and current_page[0] > 0:
                current_page[0] -= 1
                update_preview()

        def print_file():
            file_path = search.get().strip()+".pdf"
            subprocess.run(["rundll32.exe", "printui.dll,PrintUIEntry", "/p", file_path])

        # UI Components
        ctk.CTkLabel(frame, text="Filename:", text_color="Blue").grid(row=0, column=0, padx=5, pady=3)

        search = ctk.CTkEntry(frame, fg_color="white", text_color="black", placeholder_text="Enter PDF filename")
        search.grid(row=0, column=1, padx=3, pady=3)

        ctk.CTkButton(frame, text="Open", fg_color="#D8CBCE", text_color="black", command=open_file).grid(row=1, column=1, padx=20, pady=(3, 20))

        btn_prev = ctk.CTkButton(frame, text="<Prev. Page", fg_color="#F6C6C7", text_color="black", command=prev_page, state="disabled")
        btn_prev.grid(row=2, column=1, padx=20, pady=20)

        btn_next = ctk.CTkButton(frame, text="Next Page>", fg_color="#F6C6C7", text_color="black", command=next_page, state="disabled")
        btn_next.grid(row=3, column=1, padx=20, pady=20)

        ctk.CTkButton(frame, text="Print", fg_color="#D8CBCE", text_color="black", command=print_file).grid(row=4, column=1, padx=20, pady=20)

        print_win.grab_set()

    def doctor_entry(parent):
        doct = ctk.CTkToplevel(parent)
        doct.title("Doctor")
        doct.geometry("950x650")
        
        frame = ctk.CTkFrame(doct, fg_color="#e4d1d0")
        frame.place(relx=0, rely=0, relwidth=1, relheight=0.3)

        frame_det = ctk.CTkFrame(doct, fg_color="#ffffc2")
        frame_det.place(relx=0, rely=0.31, relwidth=1, relheight=0.69)

        ctk.CTkLabel(frame, text="Name:", text_color="black").grid(row=0, column=0, padx=10, pady=5, sticky="e")
        dctr_name = ctk.CTkEntry(frame, placeholder_text="Name", fg_color="white", text_color="blue")
        dctr_name.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(frame, text="Specialization:", text_color="black").grid(row=0, column=2, padx=10, pady=5)
        specialization = ctk.CTkEntry(frame, placeholder_text="specialization", fg_color="white", text_color="blue")
        specialization.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(frame, text="Location:", text_color="black").grid(row=0, column=4, padx=10, pady=5)
        location = ctk.CTkEntry(frame, placeholder_text="Location", fg_color="white", text_color="blue")
        location.grid(row=0, column=5, padx=5, pady=5, sticky="w")

        # Function to fetch data from the database
        def fetch_data():
            conn = sqlite3.connect('test_and_doctor.db')
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM doctor")
            rows = cursor.fetchall()
            conn.close()
            return rows

        # Function to populate the treeview with data
        def populate_treeview(treeview, data):
            treeview.delete(*treeview.get_children())
            for row in data:
                treeview.insert("", "end", values=row)

        def add_doctor():
            name = dctr_name.get()
            spec = specialization.get()
            loc = location.get()

            # Connect to SQLite database
            connect = sqlite3.connect('test_and_doctor.db')
            curs = connect.cursor()

            curs.execute("INSERT INTO doctor(name, specialization, location) VALUES (?, ?, ?)", (name, spec, loc))

            connect.commit()
            connect.close()

            messagebox.showinfo("ADDED SUCCESSFULLY!!!", "Doctor added successfully!")
            dctr_name.delete(0,ctk.END)
            specialization.delete(0,ctk.END)
            location.delete(0,ctk.END)
            
            # Refresh the treeview
            data = fetch_data()
            populate_treeview(treeview, data)

        def remove_doctor():
            name = dctr_name.get()
            spec = specialization.get()

            # Connect to SQLite database
            connect = sqlite3.connect('test_and_doctor.db')
            curs = connect.cursor()

            curs.execute("DELETE FROM doctor WHERE name=? AND specialization=?", (name, spec))

            connect.commit()
            connect.close()

            messagebox.showinfo("REMOVED SUCCESSFULLY!!!", "Doctor removed successfully!")
            dctr_name.delete(0,ctk.END)
            specialization.delete(0,ctk.END)
            location.delete(0,ctk.END)
            
            # Refresh the treeview
            data = fetch_data()
            populate_treeview(treeview, data)

        # Create the treeview
        columns = ("Name", "Specialization", "Location")
        treeview = ttk.Treeview(frame_det, columns=columns, show="headings")

        # Define the headings and column widths
        for col in columns:
            treeview.heading(col, text=col)
            treeview.column(col, width=100)

        # Fetch data from the database and populate the treeview
        data = fetch_data()
        populate_treeview(treeview, data)

        # Add the treeview to the frame with a scrollbar
        scrollbar = ttk.Scrollbar(frame_det, orient="vertical", command=treeview.yview)
        treeview.configure(yscrollcommand=scrollbar.set)
        treeview.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        ctk.CTkButton(frame, text="+Add", fg_color="blue", text_color="white", command=add_doctor).grid(row=1, column=0, padx=20, pady=(3, 20))
        ctk.CTkButton(frame, text="🗑Remove", fg_color="red", text_color="white", command=remove_doctor).grid(row=1, column=1, padx=20, pady=(3, 20))

        doct.grab_set()

    def test_entry(parent):
        test = ctk.CTkToplevel(parent)
        test.title("Test")
        test.geometry("950x650")
        frame = ctk.CTkFrame(test, fg_color="#fee145")
        frame.place(relx=0 ,rely=0, relwidth=1, relheight=1)

        ctk.CTkLabel(frame, text="Test Name:", text_color="black").grid(row=0, column=2, padx=(60,5), pady=(40,5))
        test_name = ctk.CTkEntry(frame, placeholder_text="Test Name", fg_color="white", text_color="black")
        test_name.grid(row=0, column=3, pady=(40,5))

        ctk.CTkLabel(frame, text="Price:", text_color="black").grid(row=0, column=4, padx=(60,5), pady=(40,5),sticky="e")
        price_entry = ctk.CTkEntry(frame, 50, placeholder_text="Price", fg_color="white", text_color="black")
        price_entry.grid(row=0, column=5, pady=(40,5), sticky="w")

        ctk.CTkLabel(frame, text="Type", text_color="black").grid(row=1, column=0, padx=20, pady=(20,5))
        test_type = ctk.CTkEntry(frame, placeholder_text="Test Type", width=200, fg_color="white", text_color="black")
        test_type.grid(row=1, column=1, pady=(20,5))

        ctk.CTkLabel(frame, text="Value", text_color="black").grid(row=1, column=3, padx=5, pady=(20,5),sticky="e")
        test_value = ctk.CTkEntry(frame, placeholder_text="Value", fg_color="white", text_color="black")
        test_value.grid(row=1, column=4, padx=0, pady=(20,5), sticky="w")

        ctk.CTkLabel(frame, text="Unit:", text_color="black").grid(row=1, column=5, padx=5, pady=(20,5))
        test_unit = ctk.CTkEntry(frame, 50, placeholder_text="Unit", fg_color="white", text_color="black")
        test_unit.grid(row=1, column=6, pady=(20,5))

        def add_to_test():
            table_name = test_name.get().strip()
            tst_type = test_type.get().strip()
            tst_value = test_value.get().strip()
            tst_unit = test_unit.get().strip()
            try:
                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()

                # Check if table exists
                curs.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
                table_exists = curs.fetchone()

                if not table_exists:
                    messagebox.showerror("Error", f"Table '{table_name}' does not exist!")
                    conn.close()
                    return

                # Insert data into the existing table
                curs.execute(f"INSERT INTO \"{table_name}\" (type, value, unit) VALUES (?, ?, ?)", 
                            (tst_type, tst_value, tst_unit))
                conn.commit()
                messagebox.showinfo("Success", f"Data added to '{table_name}' successfully!")

                conn.close()
                test_name.delete(0, ctk.END)
                price_entry.delete(0, ctk.END)
                test_type.delete(0,tk.END)
                test_value.delete(0,tk.END)
                test_unit.delete(0,tk.END)


            except Exception as e:
                messagebox.showerror("Error", str(e))

        def add():
            try:
                tst_name = test_name.get().strip()
                price = price_entry.get().strip()

                # Validation: Ensure fields are not empty
                if not tst_name or not price:
                    messagebox.showerror("Error", "Please fill in all fields.")
                    return

                # Validation: Ensure price is a valid number
                try:
                    price = float(price)  # Convert to float
                except ValueError:
                    messagebox.showerror("Error", "Price must be a valid number.")
                    return

                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()

                # Insert data, handling unique constraint for 'name'
                try:
                    curs.execute("INSERT INTO test (name, price) VALUES (?, ?)", (tst_name, price))
                    curs.execute(f"""CREATE TABLE IF NOT EXISTS "{tst_name}" (type VARCHAR(255) NOT NULL, value VARCHAR(255) NOT NULL, unit VARCHAR(255) NOT NULL)""")
                    conn.commit()
                    messagebox.showinfo("Success", "Test added successfully!")
                except sqlite3.IntegrityError:
                    messagebox.showwarning("Error", "Test name already exists!")

                conn.close()
                test_name.delete(0, ctk.END)
                price_entry.delete(0, ctk.END)
                test_type.delete(0,tk.END)
                test_value.delete(0,tk.END)
                test_unit.delete(0,tk.END)


            except Exception as e:
                messagebox.showerror("Error", str(e))
        
        def remove():
            try:
                tst_name = test_name.get().strip()

                # Validation: Ensure the test name is entered
                if not tst_name:
                    messagebox.showerror("Error", "Please enter the test name to remove.")
                    return

                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()

                # Check if the test exists before attempting deletion
                curs.execute("SELECT * FROM test WHERE name = ?", (tst_name,))
                result = curs.fetchone()

                if result:
                    # Delete the test
                    curs.execute("DELETE FROM test WHERE name = ?", (tst_name,))
                    curs.execute(f"DROP TABLE {tst_name}")
                    conn.commit()
                    messagebox.showinfo("Success", f"Test '{tst_name}' removed successfully!")
                    test_name.delete(0, ctk.END)  # Clear the entry field
                else:
                    messagebox.showwarning("Error", "Test name not found!")

                conn.close()
                test_name.delete(0, ctk.END)
                price_entry.delete(0, ctk.END)
                test_type.delete(0,tk.END)
                test_value.delete(0,tk.END)
                test_unit.delete(0,tk.END)

            except Exception as e:
                messagebox.showerror("Error", str(e))

        def update():
            try:
                tst_name = test_name.get()
                price = price_entry.get()

                if tst_name == "":
                    messagebox.showerror("Enter Value", "Enter Proper Values!!!")

                conn = sqlite3.connect("test_and_doctor.db")
                curs = conn.cursor()

                if price != "":
                    curs.execute("UPDATE test SET price = ? WHERE name = ?", (price, tst_name))
                    conn.commit()
                    messagebox.showinfo("Success", "Price updated successfully!")
                
                # elif tst_name !="" and price=="":
                #     curs.execute("")

                conn.close()
                test_name.delete(0, ctk.END)
                price_entry.delete(0, ctk.END)
                

            except Exception as e:
                messagebox.showerror("Error", e)
        
        ctk.CTkButton(frame, text="+Add", fg_color="indigo", command=add).grid(row=2, column=4, padx=80, pady=(30,5))
        ctk.CTkButton(frame, text="+Add Type", fg_color="blue", command=add_to_test).grid(row=3, column=4, padx=80, pady=5)
        ctk.CTkButton(frame, text="🗑Remove", fg_color="Maroon", command=remove).grid(row=4, column=4, padx=80, pady=5)
        ctk.CTkButton(frame, text="🛠Update", fg_color="green", command=update).grid(row=5, column=4, padx=80, pady=5)

        test.grab_set()
    
    menubar = tk.Menu(index)
    master_menu = tk.Menu(menubar, tearoff=0)
    master_menu.add_command(label="New Entry", command=lambda:new_entry(index))
    master_menu.add_command(label="View Entry", command=lambda:view_entry(index))
    master_menu.add_command(label="Print", command=lambda:print_entry(index))
    master_menu.add_command(label="Doctor", command=lambda:doctor_entry(index))
    master_menu.add_command(label="Test", command=lambda:test_entry(index))
    master_menu.add_separator()
    master_menu.add_command(label="Exit", command=exit_app)

    menubar.add_cascade(label="👨‍🔬Master", menu=master_menu)
    menubar.add_cascade(label="+New Entry", command=lambda:new_entry(index))
    menubar.add_cascade(label="🔍View Entry", command=lambda:view_entry(index))
    menubar.add_cascade(label="🖨Print", command=lambda:print_entry(index))
    menubar.add_cascade(label="🥼Doctor", command=lambda:doctor_entry(index))
    menubar.add_cascade(label="📖Test", command=lambda:test_entry(index))

    menubar.add_cascade(label="🚪Exit", command=exit_app)

    index.config(menu=menubar)

# Set window size to match the screen
    # index.geometry(f"{screen_width}x{screen_height}")
    index.minsize(600,500)

    frame = ctk.CTkFrame(index, fg_color="white")
    frame.pack(pady=20, expand=True)

    
    new_btn = ctk.CTkButton(frame, text="New Entry", width=200, height=40, fg_color="cyan", text_color="black", command=lambda:new_entry(index))
    view_btn = ctk.CTkButton(frame, text="View Entries", width=200, height=40, fg_color="light green", text_color="black", command=lambda:view_entry(index))
    test_btn = ctk.CTkButton(frame, text="Test Report", width=200, height=40, fg_color="#ffd700", text_color="black", command=lambda:test_entry(index))
    doct_btn = ctk.CTkButton(frame, text="Doctor", width=200, height=40, fg_color="#B5838D", text_color="black", command=lambda:doctor_entry(index))
    print_btn = ctk.CTkButton(frame, text="Print Report", width=200, height=40, fg_color="#76C7C0", text_color="black", command=lambda:print_entry(index))
    exit_btn = ctk.CTkButton(frame, text="Exit", width=200, height=40, fg_color="#5D89BA", text_color="black", command=exit_app)

    new_btn.grid(row=0, column=0, padx=10, pady=5)
    view_btn.grid(row=1, column=0, padx=10, pady=5)
    test_btn.grid(row=2, column=0, padx=10, pady=5)
    doct_btn.grid(row=0, column=1, padx=10, pady=5)
    print_btn.grid(row=1, column=1, padx=10, pady=5)
    exit_btn.grid(row=2, column=1, padx=10, pady=5)

    # index.mainloop()
    index.protocol("WM_DELETE_WINDOW", exit_app)

# Create the main window
if __name__ == "__main__":
    root = ctk.CTk()
    root.title("Pathology Login!!!")
    root.iconbitmap("icon.png")
    root.geometry('600x300')
    root.resizable(False, False)

    # Create Frames
    frame = ctk.CTkFrame(root, fg_color="#010101")
    frame2 = ctk.CTkFrame(root, fg_color="#011111")

    frame.place(relx=0.0, rely=0, relwidth=0.5, relheight=1)  # Left half
    frame2.place(relx=0.5, rely=0, relwidth=0.5, relheight=1)  # Right half

    # ---------------------Login------------------------
    label = ctk.CTkLabel(frame2, text="Login", text_color="#00fff7", font=("Times New Roman", 30))
    label.grid(row=0, column=0, columnspan=2, sticky="news", padx=100, pady=40)

    user_label = ctk.CTkLabel(frame2, text="UserID", text_color="white", font=("Times New Roman", 16))
    user_entry = ctk.CTkEntry(frame2, font=("Times New Roman", 16))
    user_label.grid(row=1, column=0, sticky="news", pady=5)
    user_entry.grid(row=1, column=1, sticky="news", pady=5)

    pass_label = ctk.CTkLabel(frame2, text="Password", text_color="white", font=("Times New Roman", 16))
    pass_entry = ctk.CTkEntry(frame2, show="*", font=("Times New Roman", 16))
    pass_label.grid(row=2, column=0, sticky="news", pady=5)
    pass_entry.grid(row=2, column=1, sticky="news", pady=5)

    sub_button = ctk.CTkButton(frame2, text="Submit", fg_color="#00ff00", text_color="#000080", command=entry)
    sub_button.grid(row=3, column=0, columnspan=3, sticky="news", padx=100, pady=5)

    # Load Image Correctly
    try:
        img = Image.open("icon.png").convert("RGBA")  # Convert to RGBA
        image_path = ctk.CTkImage(light_image=img, size=(150, 150))
        image_label = ctk.CTkLabel(master=frame, image=image_path, text="")  # Set master
        image_label.pack(pady=50)

    except Exception as e:
        print("Error loading image:", e)

    root.mainloop()

